"""Render the policy .docx to Markdown.

The .docx is the authoritative document. This produces the diffable copy that
GitHub renders, so the two must be regenerated together.

    python scripts/docx-to-md.py <input.docx> <output.md>

Requires python-docx.
"""
import re
import sys

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

CLAUSE = re.compile(r"^(\d+\.\d+)\t(.*)$", re.DOTALL)
SUBCLAUSE = re.compile(r"^(\([a-z]+\))\t(.*)$", re.DOTALL)


def check_docx_features(doc):
    """Fail on unsupported legally significant DOCX features.

    The Markdown is treated as a faithful rendering of the authoritative .docx.
    If the .docx contains content that this converter cannot render, continuing
    would publish an incomplete policy. The operator must resolve those features
    in Word or LibreOffice before conversion.
    """
    body = doc.element.body
    blockers = []

    revisions = body.xpath(".//w:ins | .//w:del")
    if revisions:
        blockers.append(
            f"tracked changes detected ({len(revisions)} insertion/deletion element(s)); "
            "accept or reject all changes before conversion"
        )

    comments = body.xpath(".//w:commentRangeStart | .//w:commentRangeEnd | .//w:commentReference")
    if comments:
        blockers.append(
            f"comments detected ({len(comments)} marker(s)); resolve or delete comments before conversion"
        )

    if "txbxContent" in body.xml:
        blockers.append("text boxes detected; move text box content into the document body")

    if body.xpath(".//w:fldSimple | .//w:fldChar | .//w:instrText"):
        blockers.append("fields or cross-references detected; update and convert them to static text")

    rels = list(doc.part.rels.values())
    header_text = []
    footer_text = []
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header):
            header_text.extend(p.text.strip() for p in part.paragraphs if p.text.strip())
        for part in (section.footer, section.first_page_footer, section.even_page_footer):
            footer_text.extend(p.text.strip() for p in part.paragraphs if p.text.strip())
    allowed_footer = re.compile(
        r"^(Policy on the Use of Artificial Intelligence in Legal Practice|"
        r"Protocol on the Use of Artificial Intelligence in Barristers' Practice)"
        r"\s+\|\s+Version 1\.0\s+\|\s+Page\s+of$"
    )
    disallowed_headers = [t for t in header_text if t]
    disallowed_footers = [t for t in footer_text if not allowed_footer.match(t)]
    if disallowed_headers or disallowed_footers:
        blockers.append(
            f"substantive headers/footers detected ({len(disallowed_headers)} header item(s), "
            f"{len(disallowed_footers)} footer item(s)); move legally significant content into the body"
        )

    footnote_count = 0
    endnote_count = 0
    for rel in rels:
        if rel.is_external:
            continue
        if rel.reltype.endswith("/footnotes"):
            xml = rel.target_part.blob.decode("utf-8", errors="replace")
            footnote_count += len(re.findall(r"<w:footnote\b(?![^>]*w:type=)", xml))
        if rel.reltype.endswith("/endnotes"):
            xml = rel.target_part.blob.decode("utf-8", errors="replace")
            endnote_count += len(re.findall(r"<w:endnote\b(?![^>]*w:type=)", xml))
    if footnote_count:
        blockers.append(f"footnotes detected ({footnote_count}); move footnote text into the body")
    if endnote_count:
        blockers.append(f"endnotes detected ({endnote_count}); move endnote text into the body")

    if blockers:
        joined = "; ".join(blockers)
        raise ValueError(f"unsupported DOCX feature(s): {joined}")


def style_of(p):
    try:
        return p.style.name if p.style is not None else ""
    except Exception:
        return ""


def body_items(document):
    """Yield paragraphs and tables in document order."""
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def cell_text(cell):
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "<br>".join(parts).replace("|", "\\|")


def render_table(table):
    rows = [[cell_text(c) for c in row.cells] for row in table.rows]
    if not rows:
        return []
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    head, body = rows[0], rows[1:]
    out = ["| " + " | ".join(head) + " |", "|" + "---|" * width]
    out += ["| " + " | ".join(r) + " |" for r in body]
    return out + [""]


def render_paragraph(p):
    text = p.text
    if not text.strip():
        return []
    style = style_of(p)
    if style == "Heading 1":
        return ["## " + text.strip(), ""]
    if style == "Heading 2":
        return ["### " + text.strip(), ""]

    m = CLAUSE.match(text)
    if m:
        return [f"**{m.group(1)}**  {m.group(2).strip()}", ""]
    m = SUBCLAUSE.match(text)
    if m:
        return [f"- **{m.group(1)}** {m.group(2).strip()}"]

    return [text.replace("\t", " ").strip(), ""]


def main():
    src, dst = sys.argv[1], sys.argv[2]
    document = docx.Document(src)
    try:
        check_docx_features(document)
    except ValueError as exc:
        print(f"ERROR: docx-to-md: {exc}", file=sys.stderr)
        return 1

    lines = [
        "<!-- Generated from the .docx by scripts/docx-to-md.py. Do not edit by hand. -->",
        "<!-- The .docx is the authoritative document; this is a rendering of it. -->",
        "",
    ]
    for item in body_items(document):
        if isinstance(item, Table):
            lines += render_table(item)
        else:
            lines += render_paragraph(item)

    # Collapse runs of blank lines left by empty paragraphs, and close a list
    # before non-list text so the next line isn't swallowed as a continuation.
    out, blank = [], False
    for line in lines:
        if line == "":
            if blank:
                continue
            blank = True
        else:
            if out and out[-1].startswith("- ") and not line.startswith("- "):
                out.append("")
            blank = False
        out.append(line)

    with open(dst, "w", encoding="utf-8", newline="\n") as f:
        f.write("\n".join(out).strip() + "\n")

    print(f"wrote {dst}")


if __name__ == "__main__":
    sys.exit(main())
